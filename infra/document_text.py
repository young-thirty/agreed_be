"""프로젝트 자료에서 LLM에 전달할 텍스트를 추출한다.

DeepSeek API는 이미지·파일 입력을 받지 않으므로 서버가 먼저 텍스트를 만든다.
PDF의 텍스트 레이어를 우선 사용하고, 텍스트가 없는 페이지만 Tesseract OCR로
보완한다. 실패는 빈 문자열로 강등해 파일 수집 전체를 막지 않는다.
"""

from io import BytesIO
from pathlib import Path

import pymupdf
import pytesseract
from docx import Document as DocxDocument
from PIL import Image, ImageOps

MAX_INPUT_BYTES = 12 * 1024 * 1024
MAX_OUTPUT_CHARS = 30_000
MAX_PDF_PAGES = 20
MAX_OCR_PAGES = 5
MIN_PAGE_TEXT_CHARS = 40
MAX_IMAGE_EDGE = 2600
OCR_TIMEOUT_SECONDS = 8

_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}
_TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".html", ".htm"}


def _clean(text: str) -> str:
    lines = (line.strip() for line in text.replace("\x00", "").splitlines())
    return "\n".join(line for line in lines if line)[:MAX_OUTPUT_CHARS]


def _ocr_image(image: Image.Image) -> str:
    image = ImageOps.exif_transpose(image).convert("RGB")
    if max(image.size) > MAX_IMAGE_EDGE:
        image.thumbnail((MAX_IMAGE_EDGE, MAX_IMAGE_EDGE))
    return pytesseract.image_to_string(
        image,
        lang="kor+eng",
        config="--psm 6",
        timeout=OCR_TIMEOUT_SECONDS,
    )


def _extract_pdf(data: bytes) -> str:
    blocks: list[str] = []
    ocr_pages = 0
    with pymupdf.open(stream=data, filetype="pdf") as document:
        for page_number in range(min(document.page_count, MAX_PDF_PAGES)):
            page = document.load_page(page_number)
            native = _clean(page.get_text("text"))
            if len(native) >= MIN_PAGE_TEXT_CHARS:
                blocks.append(native)
                continue
            if ocr_pages >= MAX_OCR_PAGES:
                continue
            pixmap = page.get_pixmap(matrix=pymupdf.Matrix(1.8, 1.8), alpha=False)
            image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
            scanned = _clean(_ocr_image(image))
            if scanned:
                blocks.append(scanned)
            ocr_pages += 1
    return _clean("\n\n".join(blocks))


def _extract_docx(data: bytes) -> str:
    document = DocxDocument(BytesIO(data))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    return _clean("\n".join(blocks))


def extract_document_text(data: bytes, mime_type: str | None, file_name: str) -> str:
    """지원 형식의 텍스트를 반환한다. 손상·미지원 파일은 빈 문자열이다."""

    if not data or len(data) > MAX_INPUT_BYTES:
        return ""
    suffix = Path(file_name).suffix.lower()
    mime = (mime_type or "").split(";", 1)[0].strip().lower()
    try:
        if mime == "application/pdf" or suffix == ".pdf":
            return _extract_pdf(data)
        if mime.startswith("image/") or suffix in _IMAGE_SUFFIXES:
            return _clean(_ocr_image(Image.open(BytesIO(data))))
        if (
            mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or suffix == ".docx"
        ):
            return _extract_docx(data)
        if mime.startswith("text/") or suffix in _TEXT_SUFFIXES:
            return _clean(data.decode("utf-8", errors="replace"))
    except Exception:
        return ""
    return ""
