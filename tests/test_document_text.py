import unittest
from io import BytesIO
from unittest.mock import patch

import pymupdf
from docx import Document
from PIL import Image

from infra.document_text import MAX_INPUT_BYTES, extract_document_text


class DocumentTextExtractionTest(unittest.TestCase):
    def test_extracts_native_pdf_text_without_ocr(self):
        document = pymupdf.open()
        page = document.new_page()
        page.insert_text(
            (72, 72),
            "Contract scope: homepage renewal and email login implementation.",
        )
        content = document.tobytes()
        document.close()

        with patch("infra.document_text._ocr_image") as ocr:
            extracted = extract_document_text(content, "application/pdf", "contract.pdf")

        self.assertIn("Contract scope", extracted)
        ocr.assert_not_called()

    def test_extracts_docx_paragraph_and_table(self):
        output = BytesIO()
        document = Document()
        document.add_paragraph("계약 범위: 홈페이지 리뉴얼")
        row = document.add_table(rows=1, cols=2).rows[0]
        row.cells[0].text = "납기"
        row.cells[1].text = "2026-09-20"
        document.save(output)

        extracted = extract_document_text(output.getvalue(), None, "contract.docx")

        self.assertIn("홈페이지 리뉴얼", extracted)
        self.assertIn("납기 | 2026-09-20", extracted)

    def test_routes_image_to_korean_english_ocr(self):
        output = BytesIO()
        Image.new("RGB", (100, 50), "white").save(output, format="PNG")

        with patch("infra.document_text._ocr_image", return_value="계약서 contract") as ocr:
            extracted = extract_document_text(output.getvalue(), "image/png", "scan.png")

        self.assertEqual(extracted, "계약서 contract")
        ocr.assert_called_once()

    def test_failure_and_oversized_input_are_non_blocking(self):
        self.assertEqual(extract_document_text(b"broken", "application/pdf", "bad.pdf"), "")
        self.assertEqual(
            extract_document_text(b"x" * (MAX_INPUT_BYTES + 1), "text/plain", "large.txt"),
            "",
        )


if __name__ == "__main__":
    unittest.main()
